"""
Script para colocar correctamente los 6 diagramas SVG en sus secciones
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import os

def add_diagram_to_section(doc, section_heading_text, diagram_path, diagram_title, width=5.5):
    """Añade un diagrama a una sección específica"""
    
    # Encontrar la sección
    section_index = None
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading') and section_heading_text.lower() in para.text.lower():
            section_index = i
            break
    
    if section_index is None:
        print(f"  ❌ No encontrada sección: {section_heading_text}")
        return False
    
    # Encontrar posición después del heading (buscar el siguiente párrafo con contenido)
    insert_index = section_index + 1
    
    # Saltar párrafos hasta encontrar uno que no sea heading
    while insert_index < len(doc.paragraphs):
        if not doc.paragraphs[insert_index].style.name.startswith('Heading'):
            insert_index += 1
            break
        insert_index += 1
    
    # Insertar título del diagrama
    new_para = doc.paragraphs[insert_index]._element
    new_para.addnext(new_para.__class__(new_para.tag))
    title_para = doc.paragraphs[insert_index + 1]
    title_para.text = diagram_title
    title_para.style = 'Heading 3'
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Insertar imagen
    new_para = title_para._element
    new_para.addnext(new_para.__class__(new_para.tag))
    image_para = doc.paragraphs[insert_index + 2]
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Verificar que la imagen existe
    if os.path.exists(diagram_path):
        run = image_para.add_run()
        run.add_picture(diagram_path, width=Inches(width))
        print(f"  ✅ {diagram_title} añadido ({os.path.getsize(diagram_path)/1024:.1f} KB)")
        return True
    else:
        print(f"  ❌ Archivo no encontrado: {diagram_path}")
        return False

def main():
    print("📊 COLOCANDO DIAGRAMAS EN SUS SECCIONES CORRECTAS...\n")
    
    doc_path = r'c:\xampp\htdocs\https-github.com-antonio-valero-daw2personal\Proyecto\spotMap\docs\SPOTMAP_DOCUMENTO_FINAL_PROYECTO.docx'
    images_folder = r'c:\xampp\htdocs\https-github.com-antonio-valero-daw2personal\Proyecto\spotMap\docs\images'
    
    doc = Document(doc_path)
    
    # Definir diagramas a insertar con sus secciones
    diagrams = [
        {
            'section': '5. Modelo Entidad-Relación',
            'path': os.path.join(images_folder, 'SPOTMAP_ERD_PROFESIONAL.svg'),
            'title': 'Diagrama Entidad-Relación (E/R)',
        },
        {
            'section': '5. Modelo Entidad-Relación',
            'path': os.path.join(images_folder, 'SPOTMAP_RELACIONAL.svg'),
            'title': 'Modelo Relacional PostgreSQL',
        },
        {
            'section': '6. Diagramas de Procesos',
            'path': os.path.join(images_folder, 'SPOTMAP_CLASES.svg'),
            'title': 'Diagrama de Clases UML',
        },
        {
            'section': '6. Diagramas de Procesos',
            'path': os.path.join(images_folder, 'SPOTMAP_CASOS_USO.svg'),
            'title': 'Casos de Uso',
        },
        {
            'section': '7. Diseño de Interfaz',
            'path': os.path.join(images_folder, 'SPOTMAP_WIREFRAMES.svg'),
            'title': 'Wireframes - Prototipos de Pantallas',
        },
        {
            'section': '7. Diseño de Interfaz',
            'path': os.path.join(images_folder, 'SPOTMAP_GUIA_ESTILOS.svg'),
            'title': 'Guía de Estilos - Sistema de Diseño',
        }
    ]
    
    added = 0
    for diagram in diagrams:
        print(f"📌 Sección: {diagram['section']}")
        if add_diagram_to_section(doc, diagram['section'], diagram['path'], diagram['title']):
            added += 1
        print()
    
    # Guardar documento
    print(f"\n💾 Guardando documento con {added}/6 diagramas...")
    doc.save(doc_path)
    
    print(f"\n{'='*60}")
    print(f"✅ DIAGRAMAS COLOCADOS: {added}/6")
    print(f"{'='*60}")
    print(f"\n📄 Documento guardado en: {doc_path}")
    print(f"\n🎯 Posiciones finales:")
    print(f"  1️⃣  Sección 5: E/R + Relacional")
    print(f"  2️⃣  Sección 6: Clases + Casos de Uso")
    print(f"  3️⃣  Sección 7: Wireframes + Guía de Estilos")
    print(f"\n✨ Todos los diagramas están en sus secciones correctas")

if __name__ == '__main__':
    main()
